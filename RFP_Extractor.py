import re
import spacy
import pandas as pd
import streamlit as st
from pdfminer.high_level import extract_text
from docx import Document  # For Word document support

def extract_text_from_pdf(pdf_file):
    return extract_text(pdf_file)

def extract_text_from_word(word_file):
    doc = Document(word_file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def clean_text(text):
    return re.sub(r"[\x00-\x1F\x7F]", "", text)

def extract_sentences_with_keywords(text, keywords, assigned_sentences):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    keyword_pattern = re.compile(r'(?i)\b(?:' + '|'.join(keywords) + r')\b')
    matches = []

    for sentence in sentences:
        if keyword_pattern.search(sentence) and sentence.strip() not in assigned_sentences:
            matches.append(sentence.strip())
            assigned_sentences.add(sentence.strip())

    return matches if matches else ["Not Found"]

def extract_named_entities(text, nlp, label, assigned_sentences):
    doc = nlp(text)
    matches = []

    for ent in doc.ents:
        if ent.label_ == label and ent.text not in assigned_sentences:
            matches.append(ent.text)
            assigned_sentences.add(ent.text)

    return matches if matches else ["Not Found"]

def categorize_rfp(text):
    grant_keywords = ["grant", "funding", "donation", "philanthropy", "financial aid"]
    investment_keywords = ["investment", "capital", "funding", "venture", "equity"]
    assessment_keywords = ["assessment", "evaluation", "review", "impact", "audit"]
    market_research_keywords = ["market research", "consumer research", "market analysis", "industry study", "market survey"]

    if any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in grant_keywords):
        return "Grant"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in investment_keywords):
        return "Investment"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in assessment_keywords):
        return "Assessment"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in market_research_keywords):
        return "Market Research"
    else:
        return "Uncategorized"

def process_rfp(file, file_type):
    if file_type == "pdf":
        text = extract_text_from_pdf(file)
    elif file_type == "docx":
        text = extract_text_from_word(file)
    else:
        raise ValueError("Unsupported file type")

    text = clean_text(text)
    nlp = spacy.load("./en_core_web_sm")

    # Extract and store RFP category separately
    rfp_category = categorize_rfp(text)

    # Define keyword sets
    scope_keywords = ["Scope", "Description", "Objective", "Goals", "Deliverables", "Statement of Work"]
    methodology_keywords = ["Methodology", "Approach", "Strategy", "Plan", "Implementation", "Execution", "Framework", "Process", "Techniques", "Procedures"]
    eligibility_keywords = ["Eligibility", "Eligible", "Applicants", "Who can apply", "Requirements", "Qualifications", "Criteria", "Conditions", "Target Audience"]
    budget_keywords = ["Budget", "Funding", "Cost", "Financial", "Expenses", "Price", "Pricing", "Allocation", "Payment Terms"]
    deadline_keywords = ["Deadline", "Submission", "Due Date", "Timeline", "Schedule", "Important Dates"]
    selection_process_keywords = ["Selection", "Evaluation", "Criteria", "Process", "Weighting", "Judging", "Metrics", "Assessment", "Decision"]

    assigned_sentences = set()

    extracted_info = {
        "Section": [
            "Scope of Work", "Methodology", "Eligibility",
            "Budget", "Deadlines", "Selection Process"
        ],
        "Details": [
            "\n".join(extract_sentences_with_keywords(text, scope_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, methodology_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, eligibility_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, budget_keywords, assigned_sentences)),
            "\n".join(extract_named_entities(text, nlp, "DATE", assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, selection_process_keywords, assigned_sentences)),
        ]
    }

    df = pd.DataFrame(extracted_info)
    return rfp_category, df

def save_to_word(rfp_category, df, output_path):
    doc = Document()
    doc.add_heading("RFP Extracted Information", level=1)

    # Add RFP Category at the top
    doc.add_heading("RFP Category", level=2)
    doc.add_paragraph(rfp_category, style="BodyText")

    for index, row in df.iterrows():
        doc.add_heading(row["Section"], level=2)
        doc.add_paragraph(row["Details"], style="BodyText")

    doc.save(output_path)

# Streamlit App
st.title("RFP Key Information Extractor")
st.write("Upload an RFP document (PDF or Word) to extract key details such as Scope, Eligibility, Budget, Deadlines, Selection Process, Methodology, and RFP Category.")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])

if uploaded_file is not None:
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type not in ["pdf", "docx"]:
        st.error("Unsupported file type. Please upload a PDF or Word document.")
    else:
        rfp_category, df = process_rfp(uploaded_file, file_type)

        # Show category first
        st.write("### RFP Category")
        st.success(rfp_category)

        # Display extracted information
        st.write("### Extracted Information")
        st.dataframe(df)

        # Save to Word and provide download link
        output_word = "rfp_extracted_info.docx"
        save_to_word(rfp_category, df, output_word)

        with open(output_word, "rb") as f:
            st.download_button(
                "Download Extracted Info as Word",
                f,
                file_name="rfp_extracted_info.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
